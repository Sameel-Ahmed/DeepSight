import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text, style=style)
    return p

doc = Document()

# Title Page
title = doc.add_heading('DeepSight: Underwater Image Enhancement and Marine Classification', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, '\n')
add_paragraph(doc, 'Semester Project Report', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph(doc, 'Intro to Data Science', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph(doc, '\n\n')

# Group Members
p = add_paragraph(doc, 'Group Members:')
p.style.font.bold = True
add_paragraph(doc, '1. Sameel Ahmed (Enrollment: _____________)')
add_paragraph(doc, '2. Musa Salman (Enrollment: _____________)')
add_paragraph(doc, '\n')

# Links
p = add_paragraph(doc, 'Project Links:')
p.style.font.bold = True
add_paragraph(doc, 'GitHub Repository: _________________________________________')
add_paragraph(doc, 'LinkedIn Video Demo: _______________________________________')
doc.add_page_break()

# Problem Statement
add_heading(doc, '1. Problem Statement', level=1)
add_paragraph(doc, 'Underwater environments present unique and formidable challenges for computer vision systems. As light travels through water, it undergoes severe physical distortion primarily driven by two phenomena: light attenuation and scattering. Water absorbs longer wavelengths of light, such as red and orange, within the first few meters of depth. This wavelength-dependent absorption results in underwater imagery being predominantly characterized by a strong blue-green color cast, stripping the visual data of crucial color information.')
add_paragraph(doc, 'In addition to attenuation, suspended particles such as marine snow and micro-organisms scatter the light. This scattering causes a significant loss of contrast and introduces a hazy, low-visibility effect across the images. These combined optical distortions make it exceptionally difficult to perform reliable automated marine species identification directly on raw underwater footage.')

# Methodology
add_heading(doc, '2. Detailed Methodology', level=1)
add_paragraph(doc, 'To address the complex optical challenges of underwater imagery, the DeepSight project employs a hybrid, multi-stage data science pipeline. The methodology is broadly divided into two primary phases: a physics-based mathematical restoration phase and a data-driven feature engineering and classification phase. This pipeline is implemented as an end-to-end interactive Streamlit web application, allowing for seamless ingestion, processing, and evaluation of datasets.')
add_paragraph(doc, 'The first phase focuses on mitigating the effects of light attenuation and scattering. Rather than relying on generic image filters, we implement a sequence of targeted mathematical transformations. These stages are designed to sequentially correct the color balance, enhance local contrast, and recover lost details, preparing the image for the subsequent machine learning tasks.')

add_paragraph(doc, 'The enhancement pipeline consists of the following configurable stages:')
add_paragraph(doc, 'Red Channel Compensation: Artificially restores the attenuated red pixels based on the intensity of the green channel to combat wavelength absorption.', style='List Bullet')
add_paragraph(doc, 'LAB White Balance: Converts images to the LAB color space to perform white balancing on the color channels independently, removing the dominant blue-green tint.', style='List Bullet')
add_paragraph(doc, 'Gamma Correction: Applies adaptive gamma correction based on mean brightness to illuminate dark shadow regions without overexposing highlights.', style='List Bullet')
add_paragraph(doc, 'CLAHE (Contrast Limited Adaptive Histogram Equalization): Equalizes contrast locally to reveal textures while avoiding the amplification of noise.', style='List Bullet')
add_paragraph(doc, 'Unsharp Masking: Sharpens fine surface textures such as fish scales and fins by subtracting a blurred version from the original image.', style='List Bullet')
add_paragraph(doc, 'Histogram Stretching: Maximizes the dynamic range of each color channel to recover vivid colors from low-contrast images.', style='List Bullet')

p = add_paragraph(doc, '\n[PLACEHOLDER: Image showing before and after enhancement stages]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.style.font.italic = True

add_heading(doc, '2.1. Feature Engineering and Classification', level=2)
add_paragraph(doc, 'Following the image restoration phase, the pipeline transitions to isolating the subject and extracting meaningful numerical representations. The detection module isolates the marine subject using a hybrid approach. It primarily utilizes a U-2-Net AI background removal model to generate a pixel-perfect mask. In cases where the AI fails to detect a salient object, the system falls back to classical computer vision techniques, utilizing adaptive thresholding and contour detection.')
add_paragraph(doc, 'Once the subject is cropped, a comprehensive feature extraction process is applied to generate a high-dimensional mathematical representation of the image. We extract a robust 1872-dimensional feature vector to feed into our classification models.')

add_paragraph(doc, 'The feature vector comprises the following components:')
add_paragraph(doc, 'Color Statistics (12 features): Mean, standard deviation, skewness, and kurtosis across the RGB channels.', style='List Bullet')
add_paragraph(doc, 'Color Histograms (48 features): Normalized 16-bin histograms for each of the three color channels.', style='List Bullet')
add_paragraph(doc, 'Multi-Scale LBP (48 features): Local Binary Patterns computed at multiple radii to capture structural textures like scales.', style='List Bullet')
add_paragraph(doc, 'HOG Features (1764 features): Histogram of Oriented Gradients computed on a standardized 64x64 grayscale crop to capture the anatomical shape and silhouette of the species.', style='List Bullet')

p = add_paragraph(doc, '\n[PLACEHOLDER: Diagram of the 7-step pipeline architecture]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.style.font.italic = True

p = add_paragraph(doc, '\n[PLACEHOLDER: Visualization of HOG/LBP feature extraction]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.style.font.italic = True

# Results
add_heading(doc, '3. Results', level=1)
add_paragraph(doc, 'The evaluation of the DeepSight pipeline was conducted across multiple stages, assessing both the quality of the image restoration and the performance of the classification models. The enhancement phase was benchmarked using the UIEB dataset, which provides paired raw and high-quality reference images. We utilized Peak Signal-to-Noise Ratio (PSNR) and Structural Similarity Index (SSIM) to quantitatively measure the improvement in image quality after our 6-stage pipeline was applied.')
add_paragraph(doc, 'For the classification task, the Large Scale Fish Dataset was utilized. The extracted 1872-dimensional feature vectors were used to train several models, including a Random Forest classifier, Support Vector Machines (SVM), and an ensemble voting classifier. The performance was evaluated using standard metrics such as accuracy, weighted F1-score, and per-class precision and recall.')

add_paragraph(doc, 'Key observations from the evaluation include:')
add_paragraph(doc, 'The enhancement pipeline demonstrated a significant measurable improvement in PSNR across the evaluated datasets.', style='List Bullet')
add_paragraph(doc, 'Exploratory Data Analysis confirmed the severe depletion of the red channel in raw images, validating the necessity of the Red Channel Compensation stage.', style='List Bullet')
add_paragraph(doc, 'The ensemble model showed capabilities in distinguishing between marine species based on the handcrafted feature vectors, though with room for improvement.', style='List Bullet')

p = add_paragraph(doc, '\n[PLACEHOLDER: Channel distribution histogram from EDA]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.style.font.italic = True

p = add_paragraph(doc, '\n[PLACEHOLDER: Confusion Matrix chart showing model performance]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.style.font.italic = True

# Limitations
add_heading(doc, '4. Limitations', level=1)
add_paragraph(doc, 'While the DeepSight project successfully demonstrates a comprehensive pipeline for underwater image processing, several limitations were identified during its development and evaluation. These constraints primarily affect the classification performance and the robustness of the system in highly diverse environments.')
add_paragraph(doc, 'The most significant limitation currently lies in the performance of the machine learning classification model. The feature representations and the chosen classical machine learning algorithms struggle to capture the complex, non-linear patterns required for highly accurate species identification in diverse conditions.')

add_paragraph(doc, 'Specific limitations include:')
add_paragraph(doc, 'Model Underfitting and Inaccuracy: The current classification model is not sufficiently accurate and exhibits signs of underfitting. The handcrafted features (HOG, LBP, Color Histograms), while interpretable, are not expressive enough to fully capture the intricate distinguishing characteristics of various fish species, leading to suboptimal accuracy.', style='List Bullet')
add_paragraph(doc, 'Generalization: The classical machine learning approach (Random Forest/SVM) on 1872 features struggles to generalize to entirely new underwater environments or species not well-represented in the training data.', style='List Bullet')
add_paragraph(doc, 'Processing Overhead: The extensive 6-stage mathematical enhancement pipeline, coupled with dense HOG feature extraction, introduces computational overhead that may hinder real-time video stream processing on edge devices without further optimization.', style='List Bullet')

# References
add_heading(doc, '5. References', level=1)
add_paragraph(doc, '[1] Li, C., Guo, J., Porikli, F., & Pang, Y. (2019). WaterGAN: Unsupervised Generative Network for Underwater Image Restoration. (UIEB Dataset)')
add_paragraph(doc, '[2] Ulucan, O., Karakaya, D., & Turkan, M. (2020). A Large-Scale Dataset for Fish Segmentation and Classification. (Large Scale Fish Dataset)')

doc.save("DeepSight_Project_Report.docx")
